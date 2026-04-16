"""Document loader — reads PDF, DOCX, TXT, HTML, and Markdown into raw text.

Supports local filesystem in demo mode; swap adapter for HDFS/S3 in production.
Run as a module to trigger ingestion: python -m src.retrieval.document_loader
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path

from src.config.settings import settings
from src.config.logging import get_logger
from src.connectors.files_adapter import FilesAdapter
from src.utils.ids import new_id

logger = get_logger(__name__)


@dataclass
class RawDocument:
    doc_id: str
    title: str
    source_path: str
    text: str
    file_type: str
    domain: str                # e.g. "banking" | "telco" | "government" | "general"
    ingest_timestamp: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


def _load_pdf(data: bytes, path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        content = page.extract_text() or ""
        if content.strip():
            pages.append(f"[Halaman {i}]\n{content}")
    return "\n\n".join(pages)


def _load_docx(data: bytes, _path: Path) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _load_text(data: bytes, _path: Path) -> str:
    return data.decode("utf-8", errors="replace")


def _load_html(data: bytes, _path: Path) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "html.parser")
    return soup.get_text(separator="\n")


def _load_markdown(data: bytes, _path: Path) -> str:
    return data.decode("utf-8", errors="replace")


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_text,
    ".html": _load_html,
    ".md": _load_markdown,
}


_KNOWN_DOMAINS = {"banking", "telco", "government"}


def _infer_domain(path: Path, base_path: Path) -> str:
    """Infer domain from the immediate subdirectory under base_path.

    data/sample_docs/banking/file.txt  → "banking"
    data/sample_docs/telco/file.txt    → "telco"
    data/sample_docs/government/file.txt → "government"
    data/sample_docs/file.txt          → "general"
    """
    try:
        relative = path.relative_to(base_path)
        parts = relative.parts
        if len(parts) >= 2 and parts[0] in _KNOWN_DOMAINS:
            return parts[0]
    except ValueError:
        pass
    return "banking"   # root-level files default to banking (legacy location)


def load_documents() -> list[RawDocument]:
    """Load all documents from configured source path, tagging each with its domain."""
    adapter = FilesAdapter(settings.docs_source_path)
    base_path = Path(settings.docs_source_path)
    paths = adapter.list_documents()
    documents: list[RawDocument] = []

    for path in paths:
        ext = path.suffix.lower()
        loader = _LOADERS.get(ext)
        if not loader:
            logger.warning("No loader for extension %s — skipping %s", ext, path.name)
            continue
        try:
            data = adapter.read_bytes(path)
            text = loader(data, path)
            domain = _infer_domain(path, base_path)
            documents.append(
                RawDocument(
                    doc_id=new_id(),
                    title=path.stem.replace("_", " ").replace("-", " ").title(),
                    source_path=str(path),
                    text=text,
                    file_type=ext.lstrip("."),
                    domain=domain,
                )
            )
            logger.info("Loaded: %s (domain=%s, %d chars)", path.name, domain, len(text))
        except Exception as exc:
            logger.error("Failed to load %s: %s", path.name, exc)

    return documents


if __name__ == "__main__":
    from src.config.logging import setup_logging
    from src.retrieval.chunking import chunk_documents
    from src.retrieval.embeddings import get_embeddings
    from src.retrieval.vector_store import build_vector_store

    setup_logging()
    logger.info("Starting document ingestion...")
    docs = load_documents()
    if not docs:
        logger.warning("No documents found. Add files to %s", settings.docs_source_path)
    else:
        chunks = chunk_documents(docs)
        embeddings = get_embeddings()
        build_vector_store(chunks, embeddings)
        logger.info("Ingestion complete — %d chunks indexed.", len(chunks))
